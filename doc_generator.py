"""Word 文档生成模块。

基于解析后的块结构（来自 ``format_parser.parse_markdown``）和格式配置，
使用 python-docx 生成符合格式要求的 Word 文档。
"""

from __future__ import annotations

from io import BytesIO
from typing import Dict, Iterable, Literal, TypedDict, Optional, List
import logging

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.shared import Cm, Pt
from docx.oxml.ns import qn


# 常量：中文字体集合（用于判断是否为中文文档）
CHINESE_FONT_SET_LOWER = {"宋体", "黑体", "微软雅黑", "仿宋", "楷体"}

logger = logging.getLogger(__name__)


class Block(TypedDict, total=False):
    type: Literal["title", "heading1", "heading2", "body", "table"]
    text: str
    table: Optional[List[List[str]]]


def clean_text(text: str) -> str:
    """清理正文中不希望保留的标记字符，例如 Markdown 的 `*`。"""
    return text.replace("*", "").strip()


def _get_alignment(value: str) -> int:
    """将简单的对齐字符串映射为 python-docx 的对齐常量。"""
    mapping = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get(value.lower(), WD_ALIGN_PARAGRAPH.LEFT)


def _is_english_text(text: str) -> bool:
    """简单判断一段文本是否以英文为主（ASCII 字符占比高）。"""
    if not text:
        return False
    total_chars = len(text)
    if total_chars == 0:
        return False
    ascii_count = sum(1 for ch in text if ord(ch) < 128)
    return (ascii_count / total_chars) >= 0.8


def _apply_paragraph_style(p, style_cfg: Dict[str, object], block_type: str, is_english: bool = False, is_no_indent: bool = False) -> None:
    """根据配置为段落和 run 应用样式，并对正文段落应用首行缩进。
    
    支持的 block_type: title, heading1, heading2, body
    每种类型都会应用对应的样式配置（字体、字号、加粗、对齐等）。
    """
    # 段落对齐
    alignment = style_cfg.get("alignment", "left")
    p.alignment = _get_alignment(str(alignment))

    # 行距（仅 body 需要，一般通过传入 style_cfg 时已包含）
    # 支持两种配置方式：
    # - 作为倍数（例如 1.25、1.5），会按字号 * 倍数转换为 pt
    # - 作为绝对 pt 值（例如 18），直接以 pt 应用
    line_spacing_cfg = style_cfg.get("line_spacing")
    size_pt_cfg = style_cfg.get("size_pt") or 12
    # 支持英文专用字号配置 `size_pt_en`，在 is_english 情况下优先使用；默认英文正文 11pt
    size_pt_en_cfg = style_cfg.get("size_pt_en")
    try:
        size_pt_effective = float(size_pt_en_cfg) if (is_english and size_pt_en_cfg is not None) else float(size_pt_cfg)
    except (TypeError, ValueError):
        try:
            size_pt_effective = float(size_pt_cfg)
        except (TypeError, ValueError):
            size_pt_effective = 12.0
    try:
        if block_type == "body" and line_spacing_cfg is not None:
            ls_f = float(line_spacing_cfg)
            if ls_f > 4:  # 认为这是一个绝对的 pt 值
                p.paragraph_format.line_spacing = Pt(ls_f)
            else:
                # 将倍数转换为 pt（字号 * 倍数）
                p.paragraph_format.line_spacing = Pt(float(size_pt_effective) * ls_f)
    except (TypeError, ValueError):
        # 忽略不合法的行距配置
        pass
    else:
        # 如果没有显式配置行距，对于英文正文使用 1.5 倍行距作为默认
        if block_type == "body" and is_english and line_spacing_cfg is None:
            try:
                p.paragraph_format.line_spacing = Pt(float(size_pt_effective) * 1.5)
            except (TypeError, ValueError):
                pass

    # 正文首行缩进（以"字符数"估算成厘米）
    # 注意：只有 body 类型需要首行缩进，title/heading1/heading2 不需要
    if block_type == "body":
        # 关键修复：不使用默认值2，应该从配置中获取或根据字体判断
        indent_chars = style_cfg.get("first_line_chars")  # 移除默认值2
        size_pt = style_cfg.get("size_pt")
        
        # 调试：显示实际读取的值（仅第一个body段落）
        if not hasattr(_apply_paragraph_style, "_debug_logged"):
            font_cn_dbg = style_cfg.get("font_cn", "未设置")
            logger.debug(
                "🔍 调试 - body配置: first_line_chars=%s, font_cn=%s, size_pt=%s",
                indent_chars,
                font_cn_dbg,
                size_pt,
            )
            _apply_paragraph_style._debug_logged = True
        
        try:
            indent_chars_f = float(indent_chars) if indent_chars is not None else None
            # 使用针对语言的实际字号来计算缩进
            size_pt_f = float(size_pt_effective) if size_pt_effective is not None else 0.0
            
            # 区分不同情景：
            # - 0：如果明确设置为0且是英文文档，会被统一调整为4.5字符
            # - > 0：应用缩进（中文2字符，英文统一4.5字符）
            # - None：根据字体判断默认值（中文字体→2，英文字体→4.5）
            
            # 如果该段被标记为 no_indent（例如邮件抬头），强制不缩进
            if is_no_indent:
                p.paragraph_format.first_line_indent = Cm(0)
            elif indent_chars_f == 0:
                # 如果明确设置为0，尝试判断是否为英文文档（优先用 style_cfg 提供的语言线索）
                font_cn = str(style_cfg.get("font_cn", "")).lower()
                font_en = str(style_cfg.get("font_en", "")).lower()
                is_chinese_font = font_cn in {n.lower() for n in CHINESE_FONT_SET_LOWER}
                # 当字体线索不足时（例如默认配置含中文字体），可以根据 style_cfg 中的 font_en 或者
                # 由上层传入的 flag 决定是否按英文格式处理；这里优先检测 font_en 非空且非中文。
                # 综合判断是否为英文：上层传入的 is_english 优先，其次看 font_en/ font_cn 线索
                is_english_cfg = bool(is_english) or (bool(font_en) and not is_chinese_font)

                if is_english_cfg and size_pt_f > 0:
                    # 英文文档：统一使用0.5英寸 = 1.27厘米
                    indent_cm = 1.27
                    p.paragraph_format.first_line_indent = Cm(indent_cm)
                else:
                    # 中文文档或无法判断：不缩进
                    p.paragraph_format.first_line_indent = Cm(0)
            elif indent_chars_f is not None and indent_chars_f > 0 and size_pt_f > 0:
                # 应用缩进：中文通常2字符，英文统一4.5字符（0.5英寸）
                # 检测是否为英文文档（通过字体判断）
                font_cn = str(style_cfg.get("font_cn", "")).lower()
                font_en = str(style_cfg.get("font_en", "")).lower()
                is_chinese_font = font_cn in {n.lower() for n in CHINESE_FONT_SET_LOWER}
                # 综合判断是否为英文：上层传入的 is_english 优先，其次看 font_en/ font_cn 线索
                is_english_cfg = bool(is_english) or (bool(font_en) and not is_chinese_font)

                # 对于英文文档，优先使用 0.5 英寸（1.27 cm）标准缩进；否则按字符数和字号计算
                if is_english_cfg:
                    indent_cm = 1.27
                else:
                    # 中文文档：使用字符数计算（字号 * 字符数 * pt->cm）
                    indent_cm = size_pt_f * indent_chars_f * 0.0352778
                
                p.paragraph_format.first_line_indent = Cm(indent_cm)
            elif indent_chars_f is None:
                # 配置缺失：根据字体判断默认值
                font_cn = str(style_cfg.get("font_cn", "")).lower()
                if font_cn in {n.lower() for n in CHINESE_FONT_SET_LOWER} and size_pt_f > 0:
                    # 中文字体：2字符缩进
                    indent_cm = size_pt_f * 2.0 * 0.0352778
                    p.paragraph_format.first_line_indent = Cm(indent_cm)
                elif size_pt_f > 0:
                    # 英文字体：统一使用4.5字符（0.5英寸 = 1.27厘米）
                    indent_cm = 1.27
                    p.paragraph_format.first_line_indent = Cm(indent_cm)
        except (TypeError, ValueError):
            # 如果配置异常，跳过缩进设置
            pass

    # 文字样式（在单一 run 上设置）
    run = p.runs[0] if p.runs else p.add_run()
    # 应用实际字号：英文优先使用 size_pt_en（若提供），否则使用 size_pt；若均无则用 size_pt_effective 的默认
    try:
        if size_pt_effective:
            run.font.size = Pt(float(size_pt_effective))
    except (TypeError, ValueError):
        pass

    bold = style_cfg.get("bold")
    if bold is not None:
        run.bold = bool(bold)

    # 字体：中文和英文字体的简单处理
    font_cn = style_cfg.get("font_cn") or "宋体"
    font_en = style_cfg.get("font_en") or "Times New Roman"

    # python-docx 的 font.name 主要对应西文字体
    run.font.name = str(font_en)

    # 通过底层 rFonts 同时指定东亚字体（用于中文）
    r = run._element  # noqa: SLF001  # type: ignore[attr-defined]
    rPr = r.get_or_add_rPr()
    rFonts = rPr.rFonts
    rFonts.set(qn("w:eastAsia"), str(font_cn))
    rFonts.set(qn("w:ascii"), str(font_en))
    rFonts.set(qn("w:hAnsi"), str(font_en))


def _set_page_config(doc: Document, page_cfg: Dict[str, object]) -> None:
    """根据配置设置页面大小和边距。"""
    section = doc.sections[0]

    # A4 纵向
    section.orientation = WD_ORIENT.PORTRAIT
    # A4 尺寸（单位：英寸，python-docx 内部使用 EMU）
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

    # 支持分别设置上下左右页边距，如果未指定则使用统一值
    if "margin_top_cm" in page_cfg:
        section.top_margin = Cm(float(page_cfg.get("margin_top_cm", 2.5)))
        section.bottom_margin = Cm(float(page_cfg.get("margin_bottom_cm", 2.5)))
        section.left_margin = Cm(float(page_cfg.get("margin_left_cm", 3.0)))
        section.right_margin = Cm(float(page_cfg.get("margin_right_cm", 1.5)))
    else:
        # 兼容旧配置：如果只有 margin_cm，使用统一值
        margin_cm = float(page_cfg.get("margin_cm", 2.5))
        margin = Cm(margin_cm)
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin


def generate_docx(blocks: Iterable[Block], config: Dict[str, Dict[str, object]]) -> Document:
    """根据解析结果 blocks 和配置 config 生成 Word Document 对象。

    - blocks: 由 ``format_parser.parse_markdown`` 返回的结构。
      支持的 block 类型：title, heading1, heading2, body
    - config: 通常为 ``format_parser.get_default_config()`` 的返回值。
      必须包含 page, title, heading1, heading2, body 的配置。
    """
    doc = Document()

    # 页面设置
    page_cfg = config.get("page", {})
    _set_page_config(doc, page_cfg)

    for block in blocks:
        block_type = block.get("type", "body")
        text = clean_text(block.get("text", "") or "")

        # 表格块单独处理（表格可能没有 text 内容）
        if block_type == "table":
            table_rows = block.get("table") or []
            if not table_rows:
                continue
            body_cfg = config.get("body", {})
            ncols = max(len(r) for r in table_rows)
            table = doc.add_table(rows=0, cols=ncols)
            table.style = "Table Grid"
            for r_idx, row in enumerate(table_rows):
                cells = table.add_row().cells
                for c_idx in range(ncols):
                    cell_text = clean_text(row[c_idx]) if c_idx < len(row) else ""
                    pcell = cells[c_idx].paragraphs[0]
                    run = pcell.add_run(cell_text)
                    # 字号选择：当单元格为英文时优先使用 size_pt_en
                    is_english_cell = _is_english_text(cell_text)
                    font_en = body_cfg.get("font_en") or "Times New Roman"
                    font_cn = body_cfg.get("font_cn") or "宋体"
                    size_pt_cfg = body_cfg.get("size_pt") or 12
                    size_pt_en_cfg = body_cfg.get("size_pt_en")
                    try:
                        size_pt_effective = float(size_pt_en_cfg) if (is_english_cell and size_pt_en_cfg is not None) else float(size_pt_cfg)
                    except (TypeError, ValueError):
                        size_pt_effective = float(size_pt_cfg)
                    run.font.name = str(font_en)
                    run.font.size = Pt(size_pt_effective)
                    if r_idx == 0:
                        run.bold = True
                    r = run._element  # type: ignore[attr-defined]
                    rPr = r.get_or_add_rPr()
                    rFonts = rPr.rFonts
                    rFonts.set(qn("w:eastAsia"), str(font_cn))
                    rFonts.set(qn("w:ascii"), str(font_en))
                    rFonts.set(qn("w:hAnsi"), str(font_en))
            continue

        # 非表格：正常按段落处理
        # 跳过空文本块，避免生成多余空段落
        if not text:
            continue

        # 选择对应的样式配置；支持 title, heading1, heading2, body
        # 如果 block_type 不在配置中，默认回退到 body
        style_cfg = config.get(block_type, config.get("body", {}))

        # 添加段落
        p = doc.add_paragraph()
        run = p.add_run(text)

        # 应用样式到段落与 run
        # - title/heading1/heading2: 应用字体、字号、加粗、对齐等样式
        # - body: 除了上述样式外，还会应用首行缩进和行距
        is_english = _is_english_text(text)
        is_no_indent = bool(block.get("no_indent", False))
        _apply_paragraph_style(p, style_cfg, block_type, is_english=is_english, is_no_indent=is_no_indent)

    return doc


def doc_to_bytes(doc: Document) -> bytes:
    """将 python-docx 的 Document 对象转成 bytes，供下载使用。"""
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

